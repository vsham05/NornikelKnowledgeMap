import logging
import re
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from domain.dto.document import DocumentDTO, DocumentChunkDTO
from domain.dto.image import ImageDTO
from domain.enums import DocumentType, ImageType
from ingestion.parsers.pdf_table_extract import _grid_to_markdown
from ingestion.parsers.title_slide_extract import (
    extract_authors_from_document_chunks,
    extract_authors_from_text,
    extract_all_people_from_document,
    extract_organizations_from_text,
    looks_like_organization_name,
    looks_like_person_name,
    merge_unique_names,
)

logger = logging.getLogger(__name__)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WORDS_PER_PAGE = 400
_TABLE_CAPTION_RE = re.compile(
    r"^(?:table\s+\d+|таблица\s+\d+).{0,160}$",
    re.IGNORECASE | re.UNICODE,
)


def _count_docx_words_and_breaks(doc: Document) -> tuple[int, int]:
    words = 0
    page_breaks = 0
    for para in doc.paragraphs:
        words += len(para.text.split())
        for run in para.runs:
            for br in run._element.findall(f".//{{{_W_NS}}}br"):
                if br.get(f"{{{_W_NS}}}type") == "page":
                    page_breaks += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words += len(cell.text.split())
    return words, page_breaks


def peek_docx_page_estimate(file_path: Path) -> int:
    """Estimate page count for hybrid LLM routing (DOCX has no fixed pagination)."""
    doc = Document(file_path)
    words, page_breaks = _count_docx_words_and_breaks(doc)
    if words == 0 and page_breaks == 0:
        return 0
    from_words = max(1, (words + _WORDS_PER_PAGE - 1) // _WORDS_PER_PAGE)
    from_breaks = page_breaks + 1 if page_breaks else 0
    return max(from_words, from_breaks)


def _iter_body_blocks(doc: Document):
    """Yield paragraphs and tables in document order."""
    for child in doc.element.body:
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, doc))


def _docx_table_to_block(table: Table, caption: str | None = None) -> str:
    grid = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    markdown = _grid_to_markdown(grid)
    if not markdown:
        return ""
    lines = ["[TABLE]"]
    if caption:
        lines.append(caption)
    lines.append(markdown)
    lines.append("[/TABLE]")
    return "\n".join(lines)


def _looks_like_table_caption(text: str) -> bool:
    cleaned = (text or "").strip()
    if not cleaned:
        return False
    return bool(_TABLE_CAPTION_RE.match(cleaned))


class DOCXParser:
    """Парсер DOCX документов."""

    def parse(self, file_path: Path) -> DocumentDTO:
        """Парсит DOCX файл."""
        logger.info(f"Parsing DOCX: {file_path}")

        doc = Document(file_path)

        chunks: list[DocumentChunkDTO] = []
        current_chunk_text: list[str] = []
        current_section: str | None = None
        pending_table_caption: str | None = None
        chunk_index = 0

        def flush_chunk() -> None:
            nonlocal chunk_index, current_chunk_text
            if not current_chunk_text:
                return
            chunks.append(
                DocumentChunkDTO(
                    id=uuid4(),
                    document_id=uuid4(),
                    text="\n".join(current_chunk_text),
                    chunk_index=chunk_index,
                    section_title=current_section,
                )
            )
            chunk_index += 1
            current_chunk_text = []

        def append_text(text: str) -> None:
            nonlocal pending_table_caption
            if not text:
                return
            if _looks_like_table_caption(text):
                pending_table_caption = text
                return
            current_chunk_text.append(text)
            pending_table_caption = None
            if len(current_chunk_text) >= 5:
                flush_chunk()

        def append_table(table: Table) -> None:
            nonlocal pending_table_caption
            block = _docx_table_to_block(table, pending_table_caption)
            pending_table_caption = None
            if not block:
                return
            if current_chunk_text:
                flush_chunk()
            current_chunk_text.append(block)
            flush_chunk()

        for kind, block in _iter_body_blocks(doc):
            if kind == "paragraph":
                para: Paragraph = block
                text = para.text.strip()
                if not text:
                    continue
                if para.style.name.startswith("Heading"):
                    flush_chunk()
                    current_section = text
                    continue
                append_text(text)
            else:
                append_table(block)

        flush_chunk()

        images = []
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                image_part = rel.target_part
                image_bytes = image_part.blob
                image_ext = image_part.content_type.split("/")[-1]

                image_filename = f"{file_path.stem}_{rel_id}.{image_ext}"
                image_path = f"images/{image_filename}"

                caption = self._find_caption_for_image(doc, rel_id)

                image = ImageDTO(
                    id=uuid4(),
                    document_id=uuid4(),
                    image_type=self._guess_image_type(caption),
                    file_path=image_path,
                    caption=caption,
                    ai_description="",
                )
                images.append(image)

        words, page_breaks = _count_docx_words_and_breaks(doc)
        if words == 0 and page_breaks == 0:
            estimated_pages = 0
        else:
            from_words = max(1, (words + _WORDS_PER_PAGE - 1) // _WORDS_PER_PAGE)
            from_breaks = page_breaks + 1 if page_breaks else 0
            estimated_pages = max(from_words, from_breaks)

        title = self._resolve_title(doc, file_path)
        front_matter = self._collect_front_matter(doc)
        authors = merge_unique_names(
            extract_authors_from_text(front_matter, 12_000),
            extract_authors_from_document_chunks(chunks, header_chars=2_000, max_chunks=2),
        )
        authors = extract_all_people_from_document(chunks, authors)
        organizations = extract_organizations_from_text(front_matter, 12_000)

        document = DocumentDTO(
            id=uuid4(),
            title=title,
            document_type=DocumentType.REPORT,
            file_path=str(file_path),
            chunks=chunks,
            images=images,
            authors=authors,
            organizations=organizations,
            estimated_page_count=estimated_pages or None,
        )

        for chunk in chunks:
            chunk.document_id = document.id
        for image in images:
            image.document_id = document.id

        logger.info(f"Parsed: {len(chunks)} chunks, {len(images)} images")
        return document

    @staticmethod
    def _resolve_title(doc: Document, file_path: Path) -> str:
        props = doc.core_properties
        if props and props.title:
            candidate = props.title.strip()
            if (
                candidate
                and not candidate.lower().startswith("untitled")
                and not looks_like_organization_name(candidate)
            ):
                return candidate[:300]

        heading_candidates: list[str] = []
        body_candidates: list[str] = []
        for para in doc.paragraphs[:20]:
            text = para.text.strip()
            if not text or len(text) < 8:
                continue
            if looks_like_organization_name(text) or looks_like_person_name(text):
                continue
            if para.style.name.startswith("Heading"):
                heading_candidates.append(text[:300])
                continue
            if 12 <= len(text) <= 280:
                body_candidates.append(text[:300])

        if heading_candidates:
            return heading_candidates[0]
        if body_candidates:
            return body_candidates[0]
        return file_path.stem

    @staticmethod
    def _collect_front_matter(doc: Document) -> str:
        parts: list[str] = []
        for para in doc.paragraphs[:45]:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables[:2]:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)
        return "\n".join(parts)

    def _find_caption_for_image(self, doc: Document, rel_id: str) -> str | None:
        for para in doc.paragraphs:
            text = para.text.strip()
            if any(keyword in text.lower() for keyword in ["рис.", "figure", "fig."]):
                return text
        return None

    def _guess_image_type(self, caption: str | None) -> ImageType:
        if not caption:
            return ImageType.OTHER

        caption_lower = caption.lower()

        if any(word in caption_lower for word in ["микроструктур", "microstructure", "sem", "tem"]):
            return ImageType.MICROSTRUCTURE
        if any(word in caption_lower for word in ["график", "зависимост", "plot"]):
            return ImageType.PLOT
        if any(word in caption_lower for word in ["схем", "диаграмм", "scheme"]):
            return ImageType.SCHEME

        return ImageType.OTHER
