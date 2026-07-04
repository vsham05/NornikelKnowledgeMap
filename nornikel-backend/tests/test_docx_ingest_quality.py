from docx import Document

from ingestion.parsers.docx_parser import DOCXParser, _docx_table_to_block, _looks_like_table_caption
from ingestion.parsers.title_slide_extract import looks_like_organization_name
from ingestion.upload_naming import humanize_upload_title


def test_humanize_upload_title():
    assert (
        humanize_upload_title("1_Моделирование_тектонических_нарушений.docx")
        == "1 Моделирование тектонических нарушений"
    )
    assert humanize_upload_title("report_final") == "report final"


def test_org_name_not_used_as_docx_title(tmp_path):
    path = tmp_path / "tectonic_model.docx"
    doc = Document()
    doc.add_paragraph("РАО «НОРИЛЬСКИЙ НИКЕЛЬ»")
    doc.add_heading("Моделирование тектонических нарушений", level=1)
    doc.add_paragraph("Описание верификационной модели.")
    doc.save(path)

    parsed = DOCXParser().parse(path)
    assert parsed.title == "Моделирование тектонических нарушений"
    assert looks_like_organization_name("РАО «НОРИЛЬСКИЙ НИКЕЛЬ»")


def test_docx_tables_become_structured_blocks(tmp_path):
    path = tmp_path / "with_table.docx"
    doc = Document()
    doc.add_paragraph("Таблица 1. Модули упругости")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Зона"
    table.rows[0].cells[1].text = "E, MPa"
    table.rows[0].cells[2].text = "ν"
    table.rows[1].cells[0].text = "Основная часть"
    table.rows[1].cells[1].text = "210000"
    table.rows[1].cells[2].text = "0.30"
    doc.save(path)

    parsed = DOCXParser().parse(path)
    joined = "\n".join(c.text for c in parsed.chunks)
    assert "[TABLE]" in joined
    assert "210000" in joined
    assert "0.30" in joined


def test_table_caption_detection():
    assert _looks_like_table_caption("Таблица 1. Модули упругости")
    assert not _looks_like_table_caption("Обычный абзац текста")


def test_docx_table_block_markdown():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    block = _docx_table_to_block(table, "Таблица 1")
    assert block.startswith("[TABLE]")
    assert block.endswith("[/TABLE]")
    assert "| A | B |" in block
