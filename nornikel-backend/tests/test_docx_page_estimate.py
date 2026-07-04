from pathlib import Path

from docx import Document

from ingestion.parsers.docx_parser import peek_docx_page_estimate


def test_peek_docx_page_estimate_short(tmp_path: Path):
    path = tmp_path / "short.docx"
    doc = Document()
    doc.add_paragraph(" ".join(["word"] * 200))
    doc.save(path)
    assert peek_docx_page_estimate(path) == 1


def test_peek_docx_page_estimate_long(tmp_path: Path):
    path = tmp_path / "long.docx"
    doc = Document()
    doc.add_paragraph(" ".join(["word"] * 12000))
    doc.save(path)
    assert peek_docx_page_estimate(path) >= 28
